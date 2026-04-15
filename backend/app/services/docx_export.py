from io import BytesIO
from docx import Document

from app.services.gamifyont_alignment import analyze_spec_sections


def export_spec_docx_file(spec):
    document = Document()

    document.add_heading(spec.title, 0)
    document.add_paragraph(f"Spec ID: {spec.id}")
    document.add_paragraph(f"Status: {spec.status}")

    sections = spec.sections or {}
    analysis = analyze_spec_sections(sections)
    theory = analysis.get("theory") or {}
    hexad = analysis.get("hexad") or {}
    mda = analysis.get("mda") or {}

    document.add_heading("System-assisted analysis (export snapshot)", level=1)
    document.add_paragraph(
        "Automated keyword-based summaries generated at export time. "
        "Indicative only — not a substitute for expert design or psychology review."
    )
    if theory.get("summary_en"):
        document.add_paragraph(str(theory["summary_en"]))
    if theory.get("summary_tr"):
        document.add_paragraph(str(theory["summary_tr"]))

    document.add_heading("Self-Determination Theory (SDT) tags", level=2)
    tags_en = theory.get("sdt_support_tags_en") or []
    tags_tr = theory.get("sdt_support_tags_tr") or []
    if tags_en:
        for line in tags_en:
            document.add_paragraph(line, style="List Bullet")
    else:
        document.add_paragraph("No strong SDT need signal crossed the heuristic threshold in this text.")
    if tags_tr:
        document.add_paragraph("Summary labels:")
        for line in tags_tr:
            document.add_paragraph(line, style="List Bullet")

    document.add_heading("HEXAD / MDA snapshot", level=2)
    document.add_paragraph(f"Primary HEXAD signal (guess): {hexad.get('primary_guess') or '—'}")
    buckets = (mda.get("buckets") or {}) if isinstance(mda, dict) else {}
    for name in ("mechanics", "dynamics", "aesthetics"):
        b = buckets.get(name) or {}
        sc = b.get("score")
        document.add_paragraph(f"MDA {name} coverage (0–1 heuristic): {sc if sc is not None else '—'}")
    for note in (mda.get("notes") or [])[:5]:
        document.add_paragraph(str(note), style="List Bullet")

    design = analysis.get("design_frameworks") or {}
    if design.get("copilot_hints"):
        document.add_heading("Co-pilot hints (Flow / Octalysis / Bartle heuristics)", level=2)
        for h in design["copilot_hints"][:12]:
            document.add_paragraph(str(h), style="List Bullet")
    if design.get("bartle", {}).get("primary_guess"):
        document.add_paragraph(f"Primary Bartle signal (guess): {design['bartle']['primary_guess']}")

    _six_d_labels = {
        "d1_define": "D1 — Define business objectives",
        "d2_behaviors": "D2 — Delineate target behaviors",
        "d3_players": "D3 — Describe your players",
        "d4_cycles": "D4 — Devise activity cycles",
        "d5_fun": "D5 — Don't forget the fun",
        "d6_deploy": "D6 — Deploy appropriate tools",
    }
    six_meta = sections.get("__meta::six_d")
    if six_meta and isinstance(six_meta, dict):
        document.add_heading("6D framework checklist (stored metadata)", level=2)
        document.add_paragraph(
            "Designer checklist captured in the spec (same block as Spec Studio / Export preview)."
        )
        phases = six_meta.get("phases") or {}
        order = ["d1_define", "d2_behaviors", "d3_players", "d4_cycles", "d5_fun", "d6_deploy"]
        for pid in order:
            pdata = phases.get(pid)
            if not isinstance(pdata, dict):
                continue
            label = _six_d_labels.get(pid, pid)
            done = "Done" if pdata.get("done") else "Open"
            notes = str(pdata.get("notes") or "").strip()
            line = f"{label} — {done}."
            if notes:
                line += f" Notes: {notes}"
            document.add_paragraph(line, style="List Bullet")
        for pid, pdata in phases.items():
            if pid in order or not isinstance(pdata, dict):
                continue
            done = "Done" if pdata.get("done") else "Open"
            notes = str(pdata.get("notes") or "").strip()
            line = f"{pid} — {done}."
            if notes:
                line += f" Notes: {notes}"
            document.add_paragraph(line, style="List Bullet")

    document.add_heading("Specification Content", level=1)

    for key, value in sections.items():
        if str(key).startswith("__"):
            continue
        if "::" in key:
            section_title = key.split("::", 1)[1]
        else:
            section_title = key

        document.add_heading(section_title, level=2)
        document.add_paragraph(str(value or "(empty)"))

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream
